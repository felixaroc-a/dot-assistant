"""Servicio de generación de documentos con imágenes incrustadas (P2.2).

Permite crear documentos DOCX con texto e imágenes usando python-docx.
El agente puede orquestar: generar imagen → guardar → referenciar en documento.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

from app.services.document_output_service import (
    add_docx_title_block,
    apply_docx_page_setup,
    render_markdown_lines_to_docx,
    resolve_output_path,
    sanitize_document_title,
    xlsx_header_styles,
)

log = logging.getLogger("dot.documents.images")


def _get_dot_work_dir() -> Path:
    """Compatibilidad con tests/mocks que parchean esta función."""
    from app.services.document_output_service import get_desktop_work_dir

    return get_desktop_work_dir()


def _sanitize_filename(name: str) -> str:
    return sanitize_document_title(name)


def create_docx_with_images(
    *,
    title: str,
    content: str,
    image_paths: list[str] | None = None,
    folder: str | None = None,
) -> dict:
    """Crea un DOCX con texto e imágenes incrustadas."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    filepath = resolve_output_path(kind="docx", title=title, extension="docx", folder=folder)
    filename = filepath.name

    doc = Document()
    apply_docx_page_setup(doc)
    add_docx_title_block(doc, title)

    images = image_paths or []
    image_index = 0

    paragraphs = content.split("\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        if para_text.startswith("#"):
            render_markdown_lines_to_docx(doc, para_text)
            continue

        parts = re.split(r"\[IMAGE:(\d+)\]", para_text)
        if len(parts) > 1:
            p = doc.add_paragraph()
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    if part.strip():
                        p.add_run(part)
                else:
                    try:
                        idx = int(part)
                        if idx < len(images) and Path(images[idx]).exists():
                            try:
                                p.add_run().add_picture(images[idx], width=Inches(5.0))
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                image_index += 1
                            except Exception as e:
                                p.add_run(f" [Imagen no disponible: {e}]")
                        else:
                            p.add_run(f" [Imagen {idx} no encontrada]")
                    except (ValueError, IndexError):
                        p.add_run(f" [IMAGE:{part}]")
        elif para_text.startswith("- ") or para_text.startswith("* "):
            doc.add_paragraph(para_text[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", para_text):
            doc.add_paragraph(re.sub(r"^\d+\.\s*", "", para_text), style="List Number")
        else:
            doc.add_paragraph(para_text)

    for idx, img_path in enumerate(images):
        if idx >= image_index:
            try:
                if Path(img_path).exists():
                    doc.add_page_break()
                    doc.add_heading(f"Imagen {idx + 1}", level=3)
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_para.add_run().add_picture(img_path, width=Inches(5.0))
            except Exception:
                log.warning("No se pudo incrustar imagen: %s", img_path)

    doc.core_properties.title = title
    doc.core_properties.author = "DOT IA"
    doc.save(str(filepath))
    size_bytes = filepath.stat().st_size

    log.info("DOCX creado: %s (%d bytes, %d imágenes)", filename, size_bytes, image_index)

    return {
        "ok": True,
        "filename": filename,
        "path": str(filepath),
        "size_bytes": size_bytes,
        "image_count": image_index,
    }


def create_xlsx_with_charts(
    *,
    title: str,
    data_sections: list[dict],
    output_path: str | None = None,
) -> dict:
    """Crea un XLSX con múltiples hojas y gráficos con formato profesional."""
    import openpyxl
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    from openpyxl.chart.label import DataLabelList
    from openpyxl.utils import get_column_letter

    if not data_sections:
        return {"ok": False, "error": "Se requiere al menos una sección de datos."}

    if output_path:
        filepath = Path(output_path)
        filepath.parent.mkdir(parents=True, exist_ok=True)
        filename = filepath.name
    else:
        filepath = resolve_output_path(kind="xlsx", title=title, extension="xlsx")
        filename = filepath.name

    styles = xlsx_header_styles()
    wb = openpyxl.Workbook()
    first_sheet = True
    sheet_count = 0

    chart_colors = [
        "2E75B6", "E87040", "47A86C", "9B59B6", "E74C3C",
        "3498DB", "F1C40F", "1ABC9C", "E67E22", "95A5A6",
    ]

    for section in data_sections:
        section_title = str(section.get("section_title") or f"Hoja {sheet_count + 1}")
        headers = section.get("headers", [])
        rows = section.get("rows", [])
        chart_type = str(section.get("chart_type") or "bar").lower()
        chart_title = str(section.get("chart_title") or section_title)

        if not headers or not rows:
            log.warning("Sección '%s' sin headers o rows, omitiendo.", section_title)
            continue

        safe_sheet_name = sanitize_document_title(section_title)[:31] or f"Hoja{sheet_count + 1}"

        if first_sheet:
            ws = wb.active
            ws.title = safe_sheet_name
            first_sheet = False
        else:
            ws = wb.create_sheet(title=safe_sheet_name)

        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(headers), 1))
        title_cell = ws.cell(row=1, column=1, value=section_title)
        title_cell.font = styles["title_font"]
        title_cell.alignment = openpyxl.styles.Alignment(horizontal="left", vertical="center")

        from datetime import datetime

        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=max(len(headers), 1))
        sub_cell = ws.cell(row=2, column=1, value=f"DOT IA — {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        sub_cell.font = styles["subtitle_font"]

        header_row = 4
        data_start_row = header_row + 1

        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=header_row, column=col_idx, value=str(header))
            cell.font = styles["header_font"]
            cell.fill = styles["header_fill"]
            cell.alignment = styles["header_alignment"]
            cell.border = styles["thin_border"]

        for row_idx, row in enumerate(rows, data_start_row):
            for col_idx, value in enumerate(row, 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.value = value if isinstance(value, (int, float)) else (str(value) if value is not None else "")
                cell.border = styles["thin_border"]
                cell.font = openpyxl.styles.Font(name="Calibri", size=11)

        for col_idx in range(1, len(headers) + 1):
            max_length = len(str(headers[col_idx - 1]))
            for row in rows:
                val = row[col_idx - 1] if col_idx - 1 < len(row) else ""
                max_length = max(max_length, len(str(val)))
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max_length + 3, 42)

        ws.freeze_panes = ws.cell(row=data_start_row, column=1)

        valid_chart_types = {"bar", "line", "pie"}
        if chart_type in valid_chart_types and len(headers) >= 2:
            data_end_row = data_start_row + len(rows) - 1
            data_end_col = len(headers)

            if chart_type == "bar":
                chart = BarChart()
            elif chart_type == "line":
                chart = LineChart()
            else:
                chart = PieChart()

            chart.title = chart_title
            chart.style = 10

            cats = Reference(ws, min_col=1, min_row=header_row, max_row=data_end_row)
            data_ref = Reference(ws, min_col=2, max_col=data_end_col, min_row=header_row, max_row=data_end_row)

            chart.add_data(data_ref, titles_from_data=True)
            chart.set_categories(cats)

            try:
                for idx, series in enumerate(chart.series):
                    color = chart_colors[idx % len(chart_colors)]
                    series.graphicalProperties.solidFill = color
                    if chart_type == "line":
                        series.graphicalProperties.line.width = 25000
            except Exception:
                log.debug("Error aplicando colores de serie en gráfico", exc_info=True)

            if chart_type == "pie":
                try:
                    chart.dataLabels = DataLabelList()
                    chart.dataLabels.showPercent = True
                    chart.dataLabels.showCatName = True
                except Exception:
                    log.debug("Error configurando etiquetas de pie chart", exc_info=True)

            chart_row = data_end_row + 3
            ws.add_chart(chart, f"A{chart_row}")

        sheet_count += 1

    if sheet_count == 0:
        return {"ok": False, "error": "No se generaron hojas. Verifica que las secciones tengan headers y rows."}

    wb.properties.title = title
    wb.properties.creator = "DOT IA"
    wb.save(str(filepath))
    size_bytes = filepath.stat().st_size

    log.info("XLSX multi-hoja creado: %s (%d bytes, %d hojas)", filename, size_bytes, sheet_count)

    return {
        "ok": True,
        "filename": filename,
        "path": str(filepath),
        "size_bytes": size_bytes,
        "sheet_count": sheet_count,
    }
