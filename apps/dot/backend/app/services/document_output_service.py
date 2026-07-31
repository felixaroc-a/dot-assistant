"""Rutas de salida, nombres legibles y confirmaciones en español para documentos generados."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any

# Paleta profesional DOT (Word, Excel, PowerPoint, PDF)
BRAND_BLUE = "2E75B6"
BRAND_DARK = "1F4E79"
HEADER_TEXT = "FFFFFF"
BODY_GRAY = (89, 89, 89)

KIND_SUBFOLDERS: dict[str, str] = {
    "docx": "Documentos",
    "xlsx": "Hojas de calculo",
    "pptx": "Presentaciones",
    "pdf": "Documentos",
    "txt": "Textos",
}

KIND_LABELS_ES: dict[str, str] = {
    "docx": "documento Word",
    "xlsx": "hoja de cálculo",
    "pptx": "presentación PowerPoint",
    "pdf": "documento PDF",
    "txt": "archivo de texto",
}


def get_desktop_work_dir() -> Path:
    """Carpeta DOT Trabajos en el Escritorio del usuario (Windows ES/EN, OneDrive)."""
    home = Path.home()
    candidates = [
        home / "Desktop",
        home / "Escritorio",
        home / "OneDrive" / "Desktop",
        home / "OneDrive" / "Escritorio",
        home / "Documents",
        home / "Documentos",
    ]
    work_dir: Path | None = None
    for candidate in candidates:
        if candidate.exists():
            work_dir = candidate / "DOT Trabajos"
            break
    if work_dir is None:
        work_dir = home / "DOT Trabajos"

    for subfolder in ("Documentos", "Hojas de calculo", "Textos", "Presentaciones", "Otros", "Imágenes"):
        (work_dir / subfolder).mkdir(parents=True, exist_ok=True)

    return work_dir


def sanitize_document_title(name: str, *, fallback: str = "Documento DOT") -> str:
    safe = "".join(c if c.isalnum() or c in " _-().áéíóúÁÉÍÓÚñÑ" else " " for c in name)
    safe = re.sub(r"\s+", " ", safe).strip()[:100]
    return safe or fallback


def build_output_filename(title: str, extension: str) -> str:
    """Nombre legible: «Informe Ventas - 24-07-2026.docx»."""
    safe = sanitize_document_title(title)
    date_str = datetime.now().strftime("%d-%m-%Y")
    ext = extension.lstrip(".")
    return f"{safe} - {date_str}.{ext}"


def resolve_output_path(
    *,
    kind: str,
    title: str,
    extension: str,
    folder: str | None = None,
) -> Path:
    """Resuelve ruta única en Escritorio/DOT Trabajos; evita sobrescritura con (2), (3)…"""
    work_dir = get_desktop_work_dir()
    if folder:
        target = work_dir / sanitize_document_title(folder, fallback="Otros")
    else:
        target = work_dir / KIND_SUBFOLDERS.get(kind, "Otros")
    target.mkdir(parents=True, exist_ok=True)

    filename = build_output_filename(title, extension)
    filepath = target / filename
    if not filepath.exists():
        return filepath

    stem = filepath.stem
    suffix = filepath.suffix
    counter = 2
    while True:
        candidate = target / f"{stem} ({counter}){suffix}"
        if not candidate.exists():
            return candidate
        counter += 1


def format_path_for_user(path: str | Path) -> str:
    """Muestra ruta amigable con prefijo Escritorio cuando aplica."""
    resolved = Path(path).resolve()
    home = Path.home()
    for desktop in (
        home / "Desktop",
        home / "Escritorio",
        home / "OneDrive" / "Desktop",
        home / "OneDrive" / "Escritorio",
    ):
        try:
            rel = resolved.relative_to(desktop.resolve())
            return f"Escritorio/{rel.as_posix()}"
        except ValueError:
            continue
    return str(resolved)


def build_document_confirmation(
    *,
    kind: str,
    filename: str,
    path: str | Path,
    extra_lines: list[str] | None = None,
) -> str:
    """Mensaje humano en español con ruta clara para el usuario."""
    label = KIND_LABELS_ES.get(kind, "archivo")
    resolved = Path(path).resolve()
    display_path = format_path_for_user(resolved)
    lines = [
        f"Listo. Guardé tu {label} en el Escritorio:",
        display_path,
        f"Archivo: {filename}",
        f"Ruta: {resolved}",
    ]
    if extra_lines:
        lines.extend(extra_lines)
    return "\n".join(lines)


def resolve_document_path_for_send(path: str) -> Path | None:
    """Convierte rutas amigables (Escritorio/…, ~/) a Path absoluto existente."""
    raw = (path or "").strip().strip("\"'")
    if not raw:
        return None

    if raw.startswith("Escritorio/") or raw.startswith("Escritorio\\"):
        rel = raw.split("/", 1)[-1] if "/" in raw else raw.split("\\", 1)[-1]
        home = Path.home()
        for desktop in (
            home / "Desktop",
            home / "Escritorio",
            home / "OneDrive" / "Desktop",
            home / "OneDrive" / "Escritorio",
        ):
            candidate = (desktop / rel).resolve()
            if candidate.exists():
                return candidate
        return None

    expanded = Path(raw.replace("\\", "/").replace("~/", f"{Path.home()}/")).expanduser()
    try:
        resolved = expanded.resolve()
    except OSError:
        return None
    return resolved if resolved.exists() else None


def apply_docx_page_setup(doc: Any) -> None:
    """Márgenes y tipografía base estilo informe profesional."""
    from docx.shared import Inches, Pt

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def add_docx_title_block(doc: Any, title: str) -> None:
    """Portada: título centrado + fecha en gris."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generado por DOT — {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*BODY_GRAY)
    run.italic = True

    doc.add_paragraph()


def render_markdown_lines_to_docx(doc: Any, content: str) -> None:
    """Convierte markdown básico a párrafos Word con estilo consistente."""
    from docx.shared import Pt

    for raw_line in content.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            continue

        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            for run in h.runs:
                run.font.size = Pt(12)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            for run in h.runs:
                run.font.size = Pt(14)
        elif line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            for run in h.runs:
                run.font.size = Pt(16)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s*", "", line).strip(), style="List Number")
        else:
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.line_spacing = 1.15


def xlsx_header_styles() -> dict[str, Any]:
    """Estilos compartidos para cabeceras Excel."""
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    thin = Side(style="thin", color="D9D9D9")
    return {
        "header_font": Font(bold=True, color=HEADER_TEXT, size=11, name="Calibri"),
        "header_fill": PatternFill(start_color=BRAND_BLUE, end_color=BRAND_BLUE, fill_type="solid"),
        "header_alignment": Alignment(horizontal="center", vertical="center", wrap_text=True),
        "thin_border": Border(left=thin, right=thin, top=thin, bottom=thin),
        "title_font": Font(bold=True, size=14, color=BRAND_DARK, name="Calibri"),
        "subtitle_font": Font(size=9, color="808080", italic=True, name="Calibri"),
    }
